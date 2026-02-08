from __future__ import annotations

from dataclasses import dataclass
from typing import List, Optional

from docx import Document
from PyPDF2 import PdfReader


@dataclass
class ParsedDocx:
    text: str
    paragraphs: List[str]

@dataclass
class ParsedPdf:
    text: str
    pages: List[str]

def _clean_line(line: str) -> str:
    return " ".join(line.split()).strip()


def _extract_paragraph_lines(doc: Document) -> List[str]:
    lines: List[str] = []
    for p in doc.paragraphs:
        cleaned = _clean_line(p.text)
        if cleaned:
            lines.append(cleaned)
    return lines


def _extract_table_lines(doc: Document) -> List[str]:
    lines: List[str] = []

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Each cell has paragraphs too
                for p in cell.paragraphs:
                    cleaned = _clean_line(p.text)
                    if cleaned:
                        lines.append(cleaned)

    return lines


def parse_docx(file_path: str, max_chars: Optional[int] = 80_000) -> ParsedDocx:
    doc = Document(file_path)

    lines: List[str] = []
    lines.extend(_extract_paragraph_lines(doc))
    lines.extend(_extract_table_lines(doc))

    seen = set()
    deduped: List[str] = []
    for line in lines:
        if line not in seen:
            seen.add(line)
            deduped.append(line)

    text = "\n".join(deduped)

    if max_chars is not None and len(text) > max_chars:
        text = text[:max_chars].rstrip() + "\n[Truncated]"

    return ParsedDocx(text=text, paragraphs=deduped)

def parse_pdf(file_path: str, max_chars: Optional[int] = 80_000) -> ParsedPdf:
    reader = PdfReader(file_path)

    pages: List[str] = []

    for page in reader.pages:
        raw = page.extract_text() or ""
        cleaned = _clean_line(raw)
        if cleaned:
            pages.append(cleaned)

    text = "\n".join(pages)

    if max_chars is not None and len(text) > max_chars:
        text = text[:max_chars].rstrip() + "\n[Truncated]"

    return ParsedPdf(text=text, pages=pages)
